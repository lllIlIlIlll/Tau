#!/usr/bin/env python3
"""日报三件套失败兜底 (v3.2, 2026-06-20)
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
触发: daily_report_1300 prompt 任一非 0 退出时
行为:
  1. 扫 sche_tasks/done/ 找最近一份非今日的 DOCX
  2. 复制到 temp/output/daily_<TODAY>/ 加黄条 + _FALLBACK 文件名后缀
  3. 写一份 _FALLBACK 副本到 sche_tasks/done/
  4. 写 scheduler.log WARN FALLBACK 一行
  5. 写 done 报告含 fallback 三字段 (source_date/reason/path)

CLI: python daily_report_render_fallback.py --today 2026-06-20 \
       --reason "validate_fail=1" --out-dir temp/output/daily_20260620 \
       --done-dir sche_tasks/done
"""
import argparse
import os
import re
import shutil
import sys
from datetime import datetime, timezone, timedelta

_BJT = timezone(timedelta(hours=8))
_BANNER_BG = "FFE7A0"   # 警告黄
_BANNER_FG = "1A1A1A"   # 黑字
_BANNER_BORDER = "000000"  # 黑色细边框
_FILENAME_DATE_RE = re.compile(r"非传统安全领域动态日报_(\d{8})(?!\d)")

# python-docx 颜色 (按需)
try:
    from docx.shared import Pt as _Pt, RGBColor as _RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH as _WD_ALIGN
    from docx.oxml.ns import nsdecls as _nsdecls
    from docx.oxml import parse_xml as _parse_xml
    _DOCX_OK = True
except ImportError:
    _DOCX_OK = False


def _find_latest_previous_docx(done_dir: str, today: str) -> str | None:
    """扫 done_dir 下所有 非今日 日期的 DOCX, 返回 mtime 最新一份的路径。
    跳过带 _HHMM_ 时间戳后缀的副本 (如 2026-06-19_1307_daily_report_1300_xxx.docx) —
    因为它们是同一份报告的多次执行, 不算独立日期。
    """
    if not os.path.isdir(done_dir):
        return None
    today_compact = today.replace("-", "")
    candidates = []
    for fn in os.listdir(done_dir):
        if not fn.endswith(".docx"):
            continue
        # 跳过 _HHMM_ 时间戳副本: 文件名形如 YYYY-MM-DD_HHMM_*.docx
        if re.match(r"^\d{4}-\d{2}-\d{2}_\d{4}_", fn):
            continue
        m = _FILENAME_DATE_RE.search(fn)
        if not m:
            continue
        file_date = m.group(1)
        if file_date == today_compact:
            continue  # 跳过今日
        full = os.path.join(done_dir, fn)
        if not os.path.isfile(full):
            continue
        candidates.append((os.path.getmtime(full), full, file_date))
    if not candidates:
        return None
    candidates.sort(reverse=True)  # mtime 最新在前
    return candidates[0][1]


def _add_fallback_banner(doc, today: str, source_date: str, reason: str) -> None:
    """在 DOCX 顶部四层之上插警告黄段落。
    行为: 创建新段落, 设背景色/边框/字号, 插到 doc.paragraphs[0] 之前。
    """
    if not _DOCX_OK:
        raise ImportError("python-docx 不可用, 无法加黄条")

    text = (
        f"[自动生成告警] 今日({today})日报三件套生成失败"
        f"（原因：{reason}），下方内容为 {source_date} 报告，仅供参考。"
    )
    # 在文档最前插入段落
    p = doc.paragraphs[0]._element.addprevious  # 占位 (实际用 _insert_paragraph_before)
    # 改用 python-docx 标准做法: 新建段落, 手动 XML 插入到 body[0] 之前
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    body = doc.element.body
    new_p = OxmlElement("w:p")
    # 段落属性: 居中 + 底纹 + 边框
    pPr = OxmlElement("w:pPr")
    jc = OxmlElement("w:jc")
    jc.set(qn("w:val"), "center")
    pPr.append(jc)
    # 底纹
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), _BANNER_BG)
    pPr.append(shd)
    # 边框 (4 边)
    pBdr = OxmlElement("w:pBdr")
    for side in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")  # 0.5pt
        b.set(qn("w:space"), "4")
        b.set(qn("w:color"), _BANNER_BORDER)
        pBdr.append(b)
    pPr.append(pBdr)
    # 段后间距
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:before"), "60")
    spacing.set(qn("w:after"), "120")
    pPr.append(spacing)
    new_p.append(pPr)

    # 文字 run
    r = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), "Times New Roman")
    rFonts.set(qn("w:eastAsia"), "仿宋")
    rPr.append(rFonts)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), "21")  # 10.5pt = 21 half-points
    rPr.append(sz)
    bold = OxmlElement("w:b")
    rPr.append(bold)
    color = OxmlElement("w:color")
    color.set(qn("w:val"), _BANNER_FG)
    rPr.append(color)
    r.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    t.set(qn("xml:space"), "preserve")
    r.append(t)
    new_p.append(r)

    # 插入到 body 最前
    body.insert(0, new_p)


def _add_md_fallback_banner(content: str, today: str, source_date: str, reason: str) -> str:
    """MD 顶部插 <div class="fallback-banner"> 提示块"""
    banner = (
        f'\n<div class="fallback-banner" style="'
        f'background-color: #{_BANNER_BG}; color: #{_BANNER_FG}; '
        f'border: 0.5pt solid #{_BANNER_BORDER}; '
        f'padding: 8pt 12pt; margin-bottom: 12pt; '
        f'font-weight: bold; text-align: center;'
        f'">\n'
        f'[自动生成告警] 今日({today})日报三件套生成失败'
        f'（原因：{reason}），下方内容为 {source_date} 报告，仅供参考。\n'
        f'</div>\n\n'
    )
    # 插到第一行空行之后 (即"■ ...整编■" 之后)
    lines = content.split("\n", 4)
    if len(lines) < 5:
        return banner + content
    return "\n".join(lines[:4]) + "\n" + banner + lines[4]


def _add_html_fallback_banner(content: str, today: str, source_date: str, reason: str) -> str:
    """HTML 顶部插 <div class="fallback-banner"> 提示块 (插到 <body> 之后)"""
    banner = (
        f'<div class="fallback-banner" style="'
        f'background-color: #{_BANNER_BG}; color: #{_BANNER_FG}; '
        f'border: 0.5pt solid #{_BANNER_BORDER}; '
        f'padding: 8pt 12pt; margin-bottom: 12pt; '
        f'font-weight: bold; text-align: center;'
        f'">\n'
        f'[自动生成告警] 今日({today})日报三件套生成失败'
        f'（原因：{reason}），下方内容为 {source_date} 报告，仅供参考。\n'
        f'</div>\n'
    )
    if "<body>" in content:
        return content.replace("<body>", "<body>\n" + banner, 1)
    return banner + content


def fallback_today(today: str, reason: str,
                   done_dir: str = "sche_tasks/done",
                   out_dir: str | None = None) -> dict:
    """主入口: 生成今日 fallback 三件套 + 元数据
    返回: {path_docx, path_md, path_html, source_date, source_path}
    找不到源 -> 抛 FileNotFoundError
    """
    if out_dir is None:
        out_dir = f"temp/output/daily_{today.replace('-', '')}"

    # 1. 找最近一份非今日 DOCX
    source_path = _find_latest_previous_docx(done_dir, today)
    if not source_path:
        raise FileNotFoundError(
            f"在 {done_dir} 下找不到任何非今日 ({today}) 的 DOCX 源, fallback 失败"
        )
    source_date = _FILENAME_DATE_RE.search(os.path.basename(source_path)).group(1)
    # 转回 YYYY-MM-DD
    source_date_fmt = f"{source_date[:4]}-{source_date[4:6]}-{source_date[6:8]}"

    # 2. 复制 DOCX -> 加黄条
    os.makedirs(out_dir, exist_ok=True)
    base_name = f"非传统安全领域动态日报_{today.replace('-', '')}_FALLBACK"
    out_docx = os.path.join(out_dir, f"{base_name}.docx")

    if not _DOCX_OK:
        # 退化: 直接复制不加分隔
        shutil.copy2(source_path, out_docx)
    else:
        from docx import Document
        doc = Document(source_path)
        _add_fallback_banner(doc, today, source_date_fmt, reason)
        doc.save(out_docx)

    # 3. 复制 MD / HTML (如果存在), 加 banner
    out_md = out_html = None
    md_source = os.path.join(done_dir, f"非传统安全领域动态日报_{source_date}.md")
    if os.path.exists(md_source):
        with open(md_source, encoding="utf-8") as f:
            content = f.read()
        new_content = _add_md_fallback_banner(content, today, source_date_fmt, reason)
        out_md = os.path.join(out_dir, f"{base_name}.md")
        with open(out_md, "w", encoding="utf-8") as f:
            f.write(new_content)

    html_source = os.path.join(done_dir, f"非传统安全领域动态日报_{source_date}.html")
    if os.path.exists(html_source):
        with open(html_source, encoding="utf-8") as f:
            content = f.read()
        new_content = _add_html_fallback_banner(content, today, source_date_fmt, reason)
        out_html = os.path.join(out_dir, f"{base_name}.html")
        with open(out_html, "w", encoding="utf-8") as f:
            f.write(new_content)

    # 4. 写一份 _FALLBACK 副本到 done/ (便于检索)
    done_fallback = os.path.join(done_dir, f"{base_name}.docx")
    shutil.copy2(out_docx, done_fallback)

    # 5. 写 scheduler.log
    log_path = os.path.join(done_dir, "..", "scheduler.log")
    log_path = os.path.normpath(log_path)
    try:
        os.makedirs(os.path.dirname(log_path), exist_ok=True)
        with open(log_path, "a", encoding="utf-8") as f:
            ts = datetime.now(_BJT).strftime("%Y-%m-%d %H:%M:%S")
            f.write(
                f"{ts} WARN FALLBACK daily_report_1300: {today} "
                f"reason={reason} source={source_date_fmt} "
                f"out_docx={out_docx}\n"
            )
    except OSError:
        pass  # 日志写不进不影响 fallback 主流程

    return {
        "source_date": source_date_fmt,
        "source_path": source_path,
        "path_docx": out_docx,
        "path_md": out_md,
        "path_html": out_html,
        "done_fallback": done_fallback,
    }


def main():
    parser = argparse.ArgumentParser(description="日报三件套失败兜底")
    parser.add_argument("--today", required=True, help="今日 YYYY-MM-DD (北京时间)")
    parser.add_argument("--reason", required=True, help="失败原因 (e.g. validate_fail=1)")
    parser.add_argument("--done-dir", default="sche_tasks/done", help="done 目录")
    parser.add_argument("--out-dir", default=None, help="输出目录 (默认 temp/output/daily_<YYYYMMDD>)")
    args = parser.parse_args()

    try:
        result = fallback_today(args.today, args.reason, args.done_dir, args.out_dir)
    except FileNotFoundError as e:
        print(f"[FATAL_FALLBACK] {e}", file=sys.stderr)
        sys.exit(5)

    print(f"[FALLBACK] today={args.today} source={result['source_date']} reason={args.reason}")
    print(f"  docx: {result['path_docx']}")
    if result.get("path_md"):
        print(f"  md:   {result['path_md']}")
    if result.get("path_html"):
        print(f"  html: {result['path_html']}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
