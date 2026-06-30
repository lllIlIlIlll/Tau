#!/usr/bin/env python3
"""非传统安全领域动态日报 · 三件套渲染器 v1.8
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
输入: report_data.json  (LLM 只输出纯数据)
输出: .md / .docx / .html 三件套

所有格式/颜色/间距 硬编码在此文件,LLM 不碰格式。
色板来源: v1.8 指令 模块 F.1
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
用法:
    python daily_report_render.py report_data.json
    python daily_report_render.py report_data.json --fmt md      # 只出 MD
    python daily_report_render.py report_data.json --fmt docx    # 只出 DOCX
    python daily_report_render.py report_data.json --fmt html    # 只出 HTML
"""

import json, sys, os, re
from datetime import datetime, date
from pathlib import Path
from typing import Optional

# 模块级 pt/cm 引用 (字体/段落 helper 使用)
try:
    from docx.shared import Pt as _Pt, Cm as _Cm
    Pt, Cm = _Pt, _Cm
except ImportError:
    Pt = lambda x: x  # noqa: E731
    Cm = lambda x: x  # noqa: E731

# ═══════════════════════════════════════════════════════════════
#  F.1 色板 · 单一真源 (Single Source of Truth)
# ═══════════════════════════════════════════════════════════════
class Color:
    """F.1 主题色板 — 全文色值仅在此定义"""
    PRIMARY   = "#1A1A1A"   # 正文 / 主标题 (不得用 #000000)
    ACCENT    = "#D97757"   # 强调橙 (不得用其他橙色)
    SECONDARY = "#C9B99A"   # 次色/米灰 (副标题/URL/页脚/分隔线)
    SHADE     = "#FBEEE6"   # 板块标题底纹 (不得省略)
    GRAY      = "#666666"   # 辅助灰 (来源行"来源:")

# python-docx 用 RGBColor
try:
    from docx.shared import RGBColor as _RC
    class DColor:
        PRIMARY   = _RC(0x1A, 0x1A, 0x1A)
        ACCENT    = _RC(0xD9, 0x77, 0x57)
        SECONDARY = _RC(0xC9, 0xB9, 0x9A)
        SHADE_BG  = _RC(0xFB, 0xEE, 0xE6)
        GRAY      = _RC(0x66, 0x66, 0x66)
except ImportError:
    DColor = None


# ═══════════════════════════════════════════════════════════════
#  F.2 字体规范 · 党政公文标准 (黑体标题 + 仿宋正文 + TNR 英文)
# ═══════════════════════════════════════════════════════════════
class FontCfg:
    """F.2 字体/字号/行距 单一真源。中文用系统黑体/仿宋，英文用 Times New Roman。

    各级字号对照:
      18pt = 小二  (报告大标题 / 一级标题 / 板块标题)
      15pt = 小三  (二级标题)
      14pt = 四号  (三级标题 / 副标题)
      12pt = 小四  (正文)
      10.5pt = 五号 (来源 / 图题)
      9pt = 小五   (页脚/角标)
    """
    CN_TITLE = "黑体"            # 标题族
    CN_BODY  = "仿宋"            # 正文族
    EN_FONT  = "Times New Roman"

    SZ_MAIN_TITLE = 36     # 报告大标题 (一号) — 用户指定 36pt
    SZ_SECTION    = 14     # 板块/一级标题 (四号) — 用户指定 黑体14pt
    SZ_SUB2       = 12     # 二级标题 (小四) — 用户指定 12pt
    SZ_SUB3       = 14     # 三级标题 (四号)
    SZ_BODY       = 12     # 正文 (小四)
    SZ_SOURCE     = 10.5   # 来源/图题 (五号)
    SZ_FOOTER     = 9      # 页脚/角标 (小五)
    SZ_BADGE      = 9      # 顶部标识
    SZ_DATE       = 14     # 日期行
    SZ_EN_SUB     = 9      # 英文副标题

    LINE_TITLE = 1.25      # 标题行距
    LINE_BODY  = 1.5       # 正文行距
    INDENT_BODY_CHARS = 2  # 正文段首缩进 (字符)


# ── 字体设置 helper ──
def _set_run_fonts(run, cn_font=FontCfg.CN_BODY, en_font=FontCfg.EN_FONT, sz_pt=None, bold=None):
    """为 run 同时设置中英文字体 (eastAsia + ascii + hAnsi + cs)。
    sz_pt/bold 可选覆盖默认。
    """
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    rPr = run._element.get_or_add_rPr()
    # 移除旧的 rFonts
    for old in rPr.findall(qn('w:rFonts')):
        rPr.remove(old)
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), en_font)
    rFonts.set(qn('w:eastAsia'), cn_font)
    rFonts.set(qn('w:hAnsi'), en_font)
    rFonts.set(qn('w:cs'), en_font)
    rPr.append(rFonts)
    if sz_pt is not None:
        run.font.size = Pt(sz_pt)
    if bold is not None:
        run.font.bold = bold


def _apply_para_format(para, line_spacing=None, space_before=None, space_after=None,
                      first_line_indent_cm=None):
    """统一应用段落格式。"""
    pf = para.paragraph_format
    if line_spacing is not None:
        from docx.enum.text import WD_LINE_SPACING
        pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        pf.line_spacing = line_spacing
    if space_before is not None:
        pf.space_before = Pt(space_before)
    if space_after is not None:
        pf.space_after = Pt(space_after)
    if first_line_indent_cm is not None:
        from docx.shared import Cm
        pf.first_line_indent = Cm(first_line_indent_cm)

# ═══════════════════════════════════════════════════════════════
#  固定板块定义 · 模块 D
# ═══════════════════════════════════════════════════════════════
SECTIONS = [
    {"code": "S1", "num": "一", "title": "非传统安全领域涉华要闻",
     "subtitle": None, "count_range": (3, 5)},
    {"code": "S2", "num": "二", "title": "各国非传统安全动向",
     "subtitle": None, "count_range": (5, 10)},
    {"code": "S3", "num": "三", "title": "其他热点或苗头性线索",
     "subtitle": None, "count_range": (0, 99), "optional": True},
    {"code": "S4", "num": "四", "title": "趋势观察",
     "subtitle": None, "count_range": (1, 1), "is_analysis": True},
    {"code": "S5", "num": "五", "title": "情报价值研判",
     "subtitle": None, "count_range": (3, 5), "is_signals": True},
]

# 星期映射
WEEKDAY_CN = ["星期一", "星期二", "星期三", "星期四", "星期五", "星期六", "星期日"]

# 顶部标识行固定文案
TOP_BADGE = "■ 国际非传统安全领域 · 每日情报整编 ■"
EN_SUBTITLE = "INTERNATIONAL NON-TRADITIONAL SECURITY BRIEFING"
DOC_TITLE = "非传统安全领域动态日报"
FOOTER_PREFIX = "非传统安全领域动态日报"
EMPTY_NOTE = "（无符合时效稿件）"
NO_DATA_DECL = "本期监测窗口内无可用素材"

# ═══════════════════════════════════════════════════════════════
#  数据加载
# ═══════════════════════════════════════════════════════════════
def load_data(json_path: str) -> dict:
    """加载 report_data.json, 返回字典。
    
    期望结构:
    {
      "date": "2026-06-04",          // YYYY-MM-DD
      "window": "2026-06-03 15:00 至 2026-06-04 15:00",  // 可选
      "sections": [
        {"code": "S1", "items": [
          {"pub_date": "6月4日", "source": "XX", "url": "https://...",
           "body": "...", "source_full_name": null}
        ]}
      ],
      "trends": "300-500字趋势分析文本",       // S4
      "signals": [                              // S5
        {"label": "信号1", "text": "100-150字"}
      ],
      "s3_hot": [...],      // 3.1 热点(可选)
      "s3_clues": [...]     // 3.2 苗头(可选)
    }
    """
    with open(json_path, 'r', encoding='utf-8') as f:
        data = json.load(f)
    return data

def parse_date(date_str: str) -> date:
    """解析 YYYY-MM-DD → date"""
    return datetime.strptime(date_str, "%Y-%m-%d").date()

# ═══════════════════════════════════════════════════════════════
#  H1 · 24h 窗口守门员 (v3.2 可靠性修复, 2026-06-20)
# ═══════════════════════════════════════════════════════════════
import re as _re
from datetime import timedelta as _timedelta, datetime as _datetime
from zoneinfo import ZoneInfo as _ZoneInfo

_BJT = _ZoneInfo("Asia/Shanghai")
_WINDOW_RE = _re.compile(
    r'(\d{4})-(\d{1,2})-(\d{1,2})\s+(\d{1,2}):(\d{1,2})\s*至\s*(\d{4})-(\d{1,2})-(\d{1,2})\s+(\d{1,2}):(\d{1,2})'
)


def _parse_window_string(window: str) -> tuple:
    """解析 'YYYY-MM-DD HH:MM 至 YYYY-MM-DD HH:MM' (北京时间) → (start_dt, end_dt)。
    解析失败抛 ValueError。"""
    m = _WINDOW_RE.search(window or "")
    if not m:
        raise ValueError(f"window 字符串无法解析: {window!r} (期望 'YYYY-MM-DD HH:MM 至 YYYY-MM-DD HH:MM')")
    y1, mo1, d1, h1, mi1, y2, mo2, d2, h2, mi2 = m.groups()
    start = _datetime(int(y1), int(mo1), int(d1), int(h1), int(mi1), tzinfo=_BJT)
    end = _datetime(int(y2), int(mo2), int(d2), int(h2), int(mi2), tzinfo=_BJT)
    return start, end


def _expected_24h_window(report_date_str: str) -> tuple:
    """SOP H1 严格 24h 区间: D-1 00:00 ~ D 18:00 (北京时间)。
    返回 (expected_start_dt, expected_end_dt)。"""
    d = parse_date(report_date_str)
    start = _datetime(d.year, d.month, d.day, 0, 0, tzinfo=_BJT) - _timedelta(days=1)
    end = _datetime(d.year, d.month, d.day, 18, 0, tzinfo=_BJT)
    return start, end


def enforce_window(data: dict) -> list:
    """v3.2 可靠性: 24h 窗口守门员。
    返回违例描述列表（空列表 = 通过）。行为:
      1. data['window'] 字符串必须等于 SOP H1 区间 D-1 00:00 ~ D 18:00 BJT
         - 不等 → 抛 ValueError
      2. 扫描 s1_items/s2_items/s3_hot/s3_clues 每条 pub_date
         - 落在窗口外 → 写入违例列表, **就地剔除该条**并打 stderr 警告
    注: S3 在调用方按 s3_hot/s3_clues 分别传, 不传 code='S3' 聚合。
    """
    if not isinstance(data, dict):
        raise ValueError(f"enforce_window: data 必须是 dict, 收到 {type(data).__name__}")
    report_date = data.get("date", "")
    window = data.get("window", "")
    if not report_date:
        raise ValueError("enforce_window: 缺少 data['date']")
    if not window:
        raise ValueError("enforce_window: 缺少 data['window']")
    try:
        actual_start, actual_end = _parse_window_string(window)
    except ValueError as e:
        raise ValueError(f"enforce_window: {e}") from None
    exp_start, exp_end = _expected_24h_window(report_date)
    if (actual_start, actual_end) != (exp_start, exp_end):
        raise ValueError(
            f"enforce_window: window 字符串与 SOP H1 24h 区间不符。"
            f"实际: {actual_start.isoformat()} ~ {actual_end.isoformat()};"
            f"期望: {exp_start.isoformat()} ~ {exp_end.isoformat()}"
        )

    violations = []
    for key in ("s1_items", "s2_items", "s3_hot", "s3_clues"):
        items = data.get(key)
        if not items:
            continue
        kept = []
        for item in items:
            pd = item.get("pub_date", "")
            m = _re.match(r'(\d{1,2})月(\d{1,2})日', pd or "")
            if not m:
                violations.append(f"{key}: 条目 pub_date 无法解析为 'X月X日': {pd!r}")
                continue
            month, day = int(m.group(1)), int(m.group(2))
            # 24h 窗口覆盖 D-1 全天 + D 日 00:00~18:00, 跨月跨年取 report_date 年月
            year = parse_date(report_date).year
            try:
                if month > parse_date(report_date).month:
                    item_dt = _datetime(year - 1, month, day, 12, 0, tzinfo=_BJT)
                else:
                    item_dt = _datetime(year, month, day, 12, 0, tzinfo=_BJT)
            except ValueError:
                violations.append(f"{key}: 条目日期非法 {pd!r}")
                continue
            if exp_start <= item_dt <= exp_end:
                kept.append(item)
            else:
                print(
                    f"[enforce_window] 剔除超窗条目: {key} pub_date={pd} "
                    f"(窗口 {exp_start.date()} ~ {exp_end.date()})",
                    file=sys.stderr,
                )
                violations.append(f"{key}: 剔除超窗条目 pub_date={pd}")
        data[key] = kept
    return violations


def format_date_cn(d: date) -> str:
    """date → YYYY年M月D日"""
    return f"{d.year}年{d.month}月{d.day}日"

def format_filename(d: date) -> str:
    """date → YYYYMMDD"""
    return d.strftime("%Y%m%d")

def weekday_cn(d: date) -> str:
    return WEEKDAY_CN[d.weekday()]

# ═══════════════════════════════════════════════════════════════
#  监测窗口行构造（兼容 data 已含/未含 时区后缀）
# ═══════════════════════════════════════════════════════════════
def _window_line(window: str) -> str:
    """统一构造监测窗口行：data已含'（北京时间）'则不再附加。"""
    if '（北京时间）' in window or '(北京时间)' in window:
        return f"监测窗口：{window}"
    return f"监测窗口：{window}（北京时间）"


# ═══════════════════════════════════════════════════════════════
#  Markdown 渲染器
# ═══════════════════════════════════════════════════════════════
def render_md(data: dict, output_dir: str = ".") -> str:
    """渲染 Markdown, 返回文件路径"""
    d = parse_date(data["date"])
    date_cn = format_date_cn(d)
    date_file = format_filename(d)
    wday = weekday_cn(d)
    
    lines = []
    
    # 第 1 层: 标识行
    lines.append(TOP_BADGE)
    lines.append("")
    # 第 2 层: 主标题
    lines.append(f"# 非传统安全领域动态日报（{date_cn}）")
    lines.append("")
    # 第 3 层: 英文副标题
    lines.append(EN_SUBTITLE)
    lines.append("")
    # 第 4 层: 日期行
    lines.append(f"**{date_cn}　{wday}**")
    lines.append("")
    lines.append("---")
    lines.append("")
    
    # 监测窗口声明(如有)
    if data.get("window"):
        lines.append(_window_line(data['window']))
        lines.append("")
    
    # 空素材声明
    all_empty = _is_all_empty(data)
    if all_empty:
        lines.append(NO_DATA_DECL)
        lines.append("")
    
    # 各板块
    for sec_def in SECTIONS:
        code = sec_def["code"]
        sec_title = f"## {sec_def['num']}、{sec_def['title']}"
        lines.append(sec_title)
        lines.append("")
        
        if sec_def.get("is_analysis"):
            # S4 趋势分析（三段式：核心态势/行为体联动/涉华影响方向）
            trends = data.get("trends", "")
            if isinstance(trends, dict):
                sub_items = [
                    ("**核心态势：**", trends.get("core_situation", "")),
                    ("**行为体联动：**", trends.get("actor_dynamics", "")),
                    ("**涉华影响方向：**", trends.get("china_impact_direction", "")),
                ]
                for label, txt in sub_items:
                    if txt:
                        lines.append(f"{label}{txt}")
                        lines.append("")
            elif trends:
                lines.append(trends)
                lines.append("")
            else:
                lines.append(EMPTY_NOTE)
                lines.append("")
            continue
        
        if sec_def.get("is_signals"):
            # S5 信号
            signals = data.get("signals", [])
            if signals:
                for i, sig in enumerate(signals, 1):
                    label = sig.get("label", f"信号{i}")
                    lines.append(f"**{label}：** {sig['text']}")
                    lines.append("")
            else:
                lines.append(EMPTY_NOTE)
                lines.append("")
            continue
        
        # S1 / S2 / S3 条目型板块
        items = _get_section_items(data, code)
        
        if code == "S3":
            # S3 分两个子段
            hot_items = data.get("s3_hot", [])
            clue_items = data.get("s3_clues", [])
            
            if not hot_items and not clue_items:
                if sec_def.get("optional"):
                    # 可省略板块: 不输出
                    lines = lines[:-2]  # 移除标题和空行
                    continue
                lines.append(EMPTY_NOTE)
                lines.append("")
                continue
            
            if hot_items:
                lines.append("### 3.1 社会/行业热点")
                lines.append("")
                for item in hot_items:
                    lines.append(_format_md_item(item))
                    lines.append("")
            
            if clue_items:
                lines.append("### 3.2 苗头性线索")
                lines.append("")
                for item in clue_items:
                    lines.append(_format_md_item(item))
                    lines.append("")
            continue
        
        # S1 / S2
        if not items:
            lines.append(EMPTY_NOTE)
            lines.append("")
            continue
        
        for item in items:
            lines.append(_format_md_item(item))
            lines.append("")
    
    # 页脚
    lines.append("---")
    lines.append("")
    lines.append("<div align=\"center\">*本报告由系统自动抓取指定数据源并整编，内容仅供决策参考。*</div>")
    
    content = "\n".join(lines)
    filename = f"非传统安全领域动态日报_{date_file}.md"
    filepath = os.path.join(output_dir, filename)
    with open(filepath, 'w', encoding='utf-8') as f:
        f.write(content)
    return filepath


def _strip_body_lead(body: str) -> str:
    """去除 body 中已含的段首前缀 (validate 要求 body 含前缀, render 重建自己的)
    支持中英文冒号/中文逗号/顿号结尾, 以及指出/评论/援引/转载 等变体。
    """
    import re as _re
    # 1) 主前缀: X月X日, 源名, 报道/表示/...:  (冒号结尾)
    body = _re.sub(r'^\d{1,2}月\d{1,2}日.{1,80}?(报道|表示|声明|发布|指出|评论|援引|转载报道)[\uff1a:]\s*', '', body)
    # 2) 变体: X月X日, 源名, 报道/...， (中文逗号结尾, render 重建后删除整段)
    body = _re.sub(r'^\d{1,2}月\d{1,2}日.{1,80}?(报道|表示|声明|发布|指出|评论|援引|转载报道)\uff0c\s*', '', body)
    return body

def _format_md_item(item: dict) -> str:
    """C.1 条目格式: 段首加粗 + 正文 + 来源行"""
    pub_date = item.get("pub_date", "")
    source = item.get("source", "")
    url = item.get("url", "")
    body = _strip_body_lead(item.get("body", ""))
    full_name = item.get("source_full_name")
    
    # 机构括注(C.1 规则2)
    source_display = source
    if full_name:
        source_display = f"{source}（{full_name}）"
    
    # 段首加粗含冒号 (C.1 规则1)
    header = f"**{pub_date}，{source_display}报道：**"
    
    # 正文 + 来源行 (C.1 规则4: 来源行紧跟正文,不空行)
    result = f"{header} {body}\n▸ 来源: <{url}>"
    
    return result


def _get_section_items(data: dict, code: str) -> list:
    """v3.2 可靠性: 扁平 schema 优先 (s1_items/s2_items/s3_hot/s3_clues)。
    兼容旧 sections[] 嵌套但打 stderr 警告 — 后续版本会移除兼容。"""
    key = code.lower() + "_items"
    if key in data:
        return data[key]
    # 兼容: 旧 sections[] 嵌套 (SOP v3.0- 时期)
    for sec in data.get("sections", []) or []:
        if sec.get("code") == code:
            print(
                f"[render] 警告: 检测到旧 sections[] 嵌套 schema (code={code}),"
                f" 已自动转换为扁平. 建议更新 LLM 输出改用 {key}.",
                file=sys.stderr,
            )
            return sec.get("items", [])
    return []


def _is_all_empty(data: dict) -> bool:
    """检查是否全部板块无内容"""
    for sec in data.get("sections", []):
        if sec.get("items"):
            return False
    if data.get("trends"):
        return False
    if data.get("signals"):
        return False
    if data.get("s3_hot") or data.get("s3_clues"):
        return False
    return True


# ═══════════════════════════════════════════════════════════════
#  DOCX 渲染器 (python-docx, 模块 F 全量实现)
# ═══════════════════════════════════════════════════════════════
def render_docx(data: dict, output_dir: str = ".") -> str:
    """渲染 DOCX, 返回文件路径。完全遵循模块 F 排版规范。"""
    try:
        from docx import Document
        from docx.shared import Pt, Cm, Inches, RGBColor, Emu
        from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
        from docx.enum.section import WD_ORIENT
        from docx.oxml.ns import qn, nsdecls
        from docx.oxml import parse_xml
    except ImportError:
        print("[ERROR] python-docx not installed. pip install python-docx", file=sys.stderr)
        return ""

    d = parse_date(data["date"])
    date_cn = format_date_cn(d)
    date_file = format_filename(d)
    wday = weekday_cn(d)

    doc = Document()

    # ── F.4 页面设置: A4, 2.5cm 边距 ──
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # ── 页脚: 分隔线 + 页码 ──
    _setup_footer(section, date_cn)

    # ══════════════════════════════════════════════════════
    #  F.2 文档顶部四层
    # ══════════════════════════════════════════════════════

    # 第 1 层: 标识行 黑体9pt 橙色加粗居中
    p1 = doc.add_paragraph()
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p1.space_before = Pt(0)
    p1.space_after = Pt(4)
    run1 = p1.add_run(TOP_BADGE)
    _set_run_fonts(run1, cn_font=FontCfg.CN_TITLE, en_font=FontCfg.EN_FONT, sz_pt=FontCfg.SZ_BADGE, bold=True)
    run1.font.color.rgb = DColor.ACCENT

    # 第 2 层: 主标题 黑体18pt(小二) 主色加粗居中
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.space_before = Pt(4)
    p2.space_after = Pt(2)
    run2 = p2.add_run(DOC_TITLE)
    _set_run_fonts(run2, cn_font=FontCfg.CN_TITLE, en_font=FontCfg.EN_FONT, sz_pt=FontCfg.SZ_MAIN_TITLE, bold=True)
    run2.font.color.rgb = DColor.PRIMARY

    # 第 3 层: 英文副标题 TNR 9pt 次色 字距加宽居中
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.space_before = Pt(0)
    p3.space_after = Pt(2)
    run3 = p3.add_run(EN_SUBTITLE)
    _set_run_fonts(run3, cn_font=FontCfg.CN_TITLE, en_font=FontCfg.EN_FONT, sz_pt=FontCfg.SZ_EN_SUB)
    run3.font.color.rgb = DColor.SECONDARY
    # 字距加宽 via XML
    rPr = run3._element.get_or_add_rPr()
    spacing_elem = parse_xml(f'<w:spacing {nsdecls("w")} w:val="40"/>')
    rPr.append(spacing_elem)

    # 第 4 层: 日期行 黑体14pt 主色加粗居中
    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p4.space_before = Pt(2)
    p4.space_after = Pt(12)
    run4 = p4.add_run(f"{date_cn}　{wday}")
    _set_run_fonts(run4, cn_font=FontCfg.CN_TITLE, en_font=FontCfg.EN_FONT, sz_pt=FontCfg.SZ_DATE, bold=True)
    run4.font.color.rgb = DColor.PRIMARY

    # 监测窗口声明(如有)
    if data.get("window"):
        pw = doc.add_paragraph()
        pw.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pw.space_after = Pt(8)
        rw = pw.add_run(_window_line(data['window']))
        _set_run_fonts(rw, cn_font=FontCfg.CN_BODY, en_font=FontCfg.EN_FONT, sz_pt=FontCfg.SZ_SOURCE)
        rw.font.color.rgb = DColor.GRAY
        # 监测窗口与S1之间补一个空行
        doc.add_paragraph()

    # 空素材声明
    all_empty = _is_all_empty(data)
    if all_empty:
        pe = doc.add_paragraph()
        pe.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pe.space_after = Pt(12)
        re = pe.add_run(NO_DATA_DECL)
        _set_run_fonts(re, cn_font=FontCfg.CN_BODY, en_font=FontCfg.EN_FONT, sz_pt=FontCfg.SZ_BODY)
        re.font.color.rgb = DColor.PRIMARY

    # ══════════════════════════════════════════════════════
    #  各板块
    # ══════════════════════════════════════════════════════
    for sec_def in SECTIONS:
        code = sec_def["code"]
        sec_num = sec_def["num"]
        sec_title = sec_def["title"]

        if code == "S3":
            hot_items = data.get("s3_hot", [])
            clue_items = data.get("s3_clues", [])
            if not hot_items and not clue_items:
                if sec_def.get("optional"):
                    continue  # S3 可省略
                _add_section_header(doc, sec_num, sec_title)
                _add_empty_note(doc)
                continue
            # S3 主标题
            _add_section_header(doc, sec_num, sec_title)
            if hot_items:
                _add_sub_header(doc, "3.1 社会/行业热点")
                for item in hot_items:
                    _add_docx_item(doc, item)
            if clue_items:
                _add_sub_header(doc, "3.2 苗头性线索")
                for item in clue_items:
                    _add_docx_item(doc, item)
            continue

        if sec_def.get("is_analysis"):
            # S4 趋势分析（三段式）
            _add_section_header(doc, sec_num, sec_title)
            trends = data.get("trends", "")
            if isinstance(trends, dict):
                sub_items = [
                    ("核心态势：", trends.get("core_situation", "")),
                    ("行为体联动：", trends.get("actor_dynamics", "")),
                    ("涉华影响方向：", trends.get("china_impact_direction", "")),
                ]
                for label, txt in sub_items:
                    if txt:
                        pt = doc.add_paragraph()
                        pt.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
                        pt.paragraph_format.line_spacing = 1.5
                        pt.paragraph_format.space_after = Pt(6)
                        pt.paragraph_format.first_line_indent = Cm(0.74)
                        # 加粗标签 (黑体)
                        r_label = pt.add_run(label)
                        _set_run_fonts(r_label, cn_font=FontCfg.CN_TITLE, en_font=FontCfg.EN_FONT, sz_pt=FontCfg.SZ_BODY, bold=True)
                        r_label.font.color.rgb = DColor.PRIMARY
                        # 正文 (仿宋)
                        r_txt = pt.add_run(txt)
                        _set_run_fonts(r_txt, cn_font=FontCfg.CN_BODY, en_font=FontCfg.EN_FONT, sz_pt=FontCfg.SZ_BODY)
                        r_txt.font.color.rgb = DColor.PRIMARY
            elif trends:
                pt = doc.add_paragraph()
                pt.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
                pt.paragraph_format.line_spacing = 1.5
                pt.paragraph_format.space_after = Pt(6)
                pt.paragraph_format.first_line_indent = Cm(0.74)
                rt = pt.add_run(trends)
                _set_run_fonts(rt, cn_font=FontCfg.CN_BODY, en_font=FontCfg.EN_FONT, sz_pt=FontCfg.SZ_BODY)
                rt.font.color.rgb = DColor.PRIMARY
            else:
                pt = doc.add_paragraph()
                pt.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
                pt.paragraph_format.line_spacing = 1.5
                pt.paragraph_format.first_line_indent = Cm(0.74)
                rt = pt.add_run(EMPTY_NOTE)
                _set_run_fonts(rt, cn_font=FontCfg.CN_BODY, en_font=FontCfg.EN_FONT, sz_pt=FontCfg.SZ_BODY)
                rt.font.color.rgb = DColor.PRIMARY
            continue

        if sec_def.get("is_signals"):
            # S5 信号
            signals = data.get("signals", [])
            if not signals:
                continue  # 信号板块无内容时省略
            _add_section_header(doc, sec_num, sec_title)
            for i, sig in enumerate(signals, 1):
                label = sig.get("label", f"信号{i}")
                ps = doc.add_paragraph()
                ps.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
                ps.paragraph_format.line_spacing = 1.5
                ps.paragraph_format.space_after = Pt(6)
                ps.paragraph_format.first_line_indent = Cm(0.74)
                # 加粗标签 (黑体)
                rl = ps.add_run(f"{label}：")
                _set_run_fonts(rl, cn_font=FontCfg.CN_TITLE, en_font=FontCfg.EN_FONT, sz_pt=FontCfg.SZ_BODY, bold=True)
                rl.font.color.rgb = DColor.ACCENT
                # 正文 (仿宋)
                rs = ps.add_run(f" {sig['text']}")
                _set_run_fonts(rs, cn_font=FontCfg.CN_BODY, en_font=FontCfg.EN_FONT, sz_pt=FontCfg.SZ_BODY)
                rs.font.color.rgb = DColor.PRIMARY
            continue

        # S1 / S2 条目型
        items = _get_section_items(data, code)
        if not items and sec_def.get("optional"):
            continue
        _add_section_header(doc, sec_num, sec_title)
        if not items:
            _add_empty_note(doc)
            continue
        for item in items:
            _add_docx_item(doc, item)

    # 保存
    filename = f"非传统安全领域动态日报_{date_file}.docx"
    filepath = os.path.join(output_dir, filename)
    doc.save(filepath)
    return filepath


def _add_section_header(doc, sec_num: str, sec_title: str):
    """F.3 板块标题: 14pt 橙色加粗 + 底纹 + 左边框4mm"""
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import nsdecls
    from docx.oxml import parse_xml

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.space_before = Pt(16)
    p.space_after = Pt(8)

    # 底纹 + 左边框 via paragraph shading
    pPr = p._element.get_or_add_pPr()
    # Shading (背景色 #FBEEE6)
    shading = parse_xml(
        f'<w:shd {nsdecls("w")} w:fill="FBEEE6" w:val="clear"/>'
    )
    pPr.append(shading)
    # 左边框 4mm
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:left w:val="single" w:sz="32" w:space="8" w:color="D97757"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)

    # 文字
    full_title = f"{sec_num}、{sec_title}"
    run = p.add_run(full_title)
    _set_run_fonts(run, cn_font=FontCfg.CN_TITLE, en_font=FontCfg.EN_FONT, sz_pt=FontCfg.SZ_SECTION, bold=True)
    run.font.color.rgb = DColor.ACCENT


def _add_sub_header(doc, text: str):
    """S3 子标题: 黑体 12pt"""
    p = doc.add_paragraph()
    p.space_before = Pt(8)
    p.space_after = Pt(4)
    run = p.add_run(text)
    _set_run_fonts(run, cn_font=FontCfg.CN_TITLE, en_font=FontCfg.EN_FONT, sz_pt=FontCfg.SZ_SUB2, bold=True)
    run.font.color.rgb = DColor.PRIMARY


def _add_docx_item(doc, item: dict):
    """渲染单条新闻: 段首加粗 + 正文 + 来源行(三色)"""
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING

    pub_date = item.get("pub_date", "")
    source = item.get("source", "")
    url = item.get("url", "")
    body = _strip_body_lead(item.get("body", ""))
    full_name = item.get("source_full_name")

    source_display = source
    if full_name:
        source_display = f"{source}（{full_name}）"

    header_text = f"{pub_date}，{source_display}报道："

    # 正文段落
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Cm(0.74)

    # 段首加粗 (黑体)
    rh = p.add_run(header_text)
    _set_run_fonts(rh, cn_font=FontCfg.CN_TITLE, en_font=FontCfg.EN_FONT, sz_pt=FontCfg.SZ_BODY, bold=True)
    rh.font.color.rgb = DColor.PRIMARY

    # 正文 (仿宋)
    rb = p.add_run(f" {body}")
    _set_run_fonts(rb, cn_font=FontCfg.CN_BODY, en_font=FontCfg.EN_FONT, sz_pt=FontCfg.SZ_BODY)
    rb.font.color.rgb = DColor.PRIMARY

    # 来源行 (三色: ▸橙 / 来源:灰 / URL次色)
    ps = doc.add_paragraph()
    ps.paragraph_format.space_before = Pt(0)
    ps.paragraph_format.space_after = Pt(6)

    ra = ps.add_run("▸ ")
    _set_run_fonts(ra, cn_font=FontCfg.CN_BODY, en_font=FontCfg.EN_FONT, sz_pt=FontCfg.SZ_SOURCE)
    ra.font.color.rgb = DColor.ACCENT

    rl = ps.add_run("来源: ")
    _set_run_fonts(rl, cn_font=FontCfg.CN_BODY, en_font=FontCfg.EN_FONT, sz_pt=FontCfg.SZ_SOURCE)
    rl.font.color.rgb = DColor.GRAY

    ru = ps.add_run(url)
    _set_run_fonts(ru, cn_font=FontCfg.CN_BODY, en_font=FontCfg.EN_FONT, sz_pt=FontCfg.SZ_SOURCE)
    ru.font.color.rgb = DColor.SECONDARY


def _add_empty_note(doc):
    """空板块提示"""
    p = doc.add_paragraph()
    p.space_after = Pt(6)
    run = p.add_run(EMPTY_NOTE)
    _set_run_fonts(run, cn_font=FontCfg.CN_BODY, en_font=FontCfg.EN_FONT, sz_pt=FontCfg.SZ_BODY)
    run.font.color.rgb = DColor.GRAY


def _setup_footer(section, date_cn: str):
    """页脚: 分隔线 + 页码, 次色"""
    from docx.oxml.ns import nsdecls
    from docx.oxml import parse_xml

    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.alignment = 1  # CENTER (WD_ALIGN_PARAGRAPH.CENTER=1, 之前误写为2=RIGHT)

    # 分隔线
    pPr = fp._element.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="4" w:space="4" w:color="C9B99A"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)

    # 页码
    run1 = fp.add_run(f"{FOOTER_PREFIX} · {date_cn}　")
    _set_run_fonts(run1, cn_font=FontCfg.CN_BODY, en_font=FontCfg.EN_FONT, sz_pt=9)
    run1.font.color.rgb = DColor.SECONDARY

    # PAGE field
    fld_xml = (
        f'<w:fldSimple {nsdecls("w")} w:instr=" PAGE \\* MERGEFORMAT ">'
        f'<w:r><w:rPr><w:color w:val="C9B99A"/></w:rPr><w:t>1</w:t></w:r>'
        f'</w:fldSimple>'
    )
    fld_elem = parse_xml(fld_xml)
    fp._element.append(fld_elem)

# ═══════════════════════════════════════════════════════════════
#  HTML 渲染器 (claude_html_theme 风格, 自包含)
# ═══════════════════════════════════════════════════════════════
def render_html(data: dict, output_dir: str = ".") -> str:
    """渲染自包含 HTML, 返回文件路径"""
    d = parse_date(data["date"])
    date_cn = format_date_cn(d)
    date_file = format_filename(d)
    wday = weekday_cn(d)

    all_empty = _is_all_empty(data)

    html_parts = []
    html_parts.append(_html_head(date_cn, wday))
    html_parts.append(_html_top_banner())
    html_parts.append(_html_title(date_cn, wday))

    if data.get("window"):
        html_parts.append('<p class="window">' + _esc(_window_line(data["window"])) + '</p>')

    if all_empty:
        html_parts.append('<p class="empty-decl">' + NO_DATA_DECL + '</p>')

    # 板块
    for sec_def in SECTIONS:
        code = sec_def["code"]
        sec_num = sec_def["num"]
        sec_title = sec_def["title"]

        html_parts.append('<div class="section">')
        html_parts.append('<h2 class="section-title">' + sec_num + '、' + _esc(sec_title) + '</h2>')

        if sec_def.get("is_analysis"):
            trends = data.get("trends", "")
            if isinstance(trends, dict):
                sub_items = [
                    ("核心态势：", trends.get("core_situation", "")),
                    ("行为体联动：", trends.get("actor_dynamics", "")),
                    ("涉华影响方向：", trends.get("china_impact_direction", "")),
                ]
                for label, txt in sub_items:
                    if txt:
                        html_parts.append('<div class="analysis-text"><strong>' + _esc(label) + '</strong>' + _esc(txt) + '</div>')
            elif trends:
                html_parts.append('<div class="analysis-text">' + _esc(trends) + '</div>')
            else:
                html_parts.append('<div class="analysis-text">' + _esc(EMPTY_NOTE) + '</div>')

        elif sec_def.get("is_signals"):
            signals = data.get("signals", [])
            if signals:
                for i, sig in enumerate(signals, 1):
                    label = sig.get("label", "信号" + str(i))
                    html_parts.append(
                        '<div class="signal-item">'
                        '<span class="signal-label">' + _esc(label) + '：</span>'
                        + _esc(sig["text"])
                        + '</div>'
                    )
            else:
                html_parts.append('<p class="empty-note">' + EMPTY_NOTE + '</p>')

        elif code == "S3":
            hot_items = data.get("s3_hot", [])
            clue_items = data.get("s3_clues", [])
            if not hot_items and not clue_items:
                if sec_def.get("optional"):
                    html_parts.pop()  # remove <div class="section">
                    html_parts.pop()  # remove <h2>
                    continue
                html_parts.append('<p class="empty-note">' + EMPTY_NOTE + '</p>')
            else:
                if hot_items:
                    html_parts.append('<h3 class="sub-title">3.1 社会/行业热点</h3>')
                    for item in hot_items:
                        html_parts.append(_html_item(item))
                if clue_items:
                    html_parts.append('<h3 class="sub-title">3.2 苗头性线索</h3>')
                    for item in clue_items:
                        html_parts.append(_html_item(item))
        else:
            # S1 / S2
            items = _get_section_items(data, code)
            if not items:
                html_parts.append('<p class="empty-note">' + EMPTY_NOTE + '</p>')
            else:
                for item in items:
                    html_parts.append(_html_item(item))

        html_parts.append('</div>')  # .section

    html_parts.append(_html_footer(date_cn))
    html_parts.append('</body></html>')

    content = "\n".join(html_parts)
    filename = "非传统安全领域动态日报_" + date_file + ".html"
    filepath = os.path.join(output_dir, filename)
    with open(filepath, 'w', encoding='utf-8') as f:
        f.write(content)
    return filepath


def _esc(text: str) -> str:
    """HTML 转义"""
    import html as _html_mod
    return _html_mod.escape(str(text), quote=True)


def _html_head(date_cn: str, wday: str) -> str:
    title_text = "非传统安全领域动态日报（" + date_cn + "）"
    return (
        '<!DOCTYPE html>\n'
        '<html lang="zh-CN">\n'
        '<head>\n'
        '<meta charset="UTF-8">\n'
        '<meta name="viewport" content="width=device-width, initial-scale=1.0">\n'
        '<title>' + _esc(title_text) + '</title>\n'
        '<style>\n'
        '  @page { size: A4; margin: 2.5cm; }\n'
        '  * { margin: 0; padding: 0; box-sizing: border-box; }\n'
        '  body {\n'
        '    font-family: "SimSun", "仿宋", "FangSong", "Times New Roman", "PingFang SC", serif;\n'
        '    color: ' + Color.PRIMARY + ';\n'
        '    max-width: 210mm;\n'
        '    margin: 0 auto;\n'
        '    padding: 2.5cm;\n'
        '    line-height: 1.5;\n'
        '    font-size: 12pt;\n'
        '    background: #fff;\n'
        '  }\n'
        '  .top-badge {\n'
        '    text-align: center;\n'
        '    font-size: 9pt;\n'
        '    font-weight: bold;\n'
        '    font-family: "SimHei", "黑体", sans-serif;\n'
        '    color: ' + Color.ACCENT + ';\n'
        '    margin-bottom: 4pt;\n'
        '  }\n'
        '  .main-title {\n'
        '    text-align: center;\n'
        '    font-size: 36pt;\n'
        '    font-weight: bold;\n'
        '    font-family: "SimHei", "黑体", "Microsoft YaHei", sans-serif;\n'
        '    color: ' + Color.PRIMARY + ';\n'
        '    margin: 4pt 0 2pt;\n'
        '  }\n'
        '  .en-subtitle {\n'
        '    text-align: center;\n'
        '    font-size: 9pt;\n'
        '    font-family: "Times New Roman", serif;\n'
        '    color: ' + Color.SECONDARY + ';\n'
        '    letter-spacing: 0.3em;\n'
        '    margin-bottom: 2pt;\n'
        '  }\n'
        '  .date-line {\n'
        '    text-align: center;\n'
        '    font-size: 14pt;\n'
        '    font-weight: bold;\n'
        '    font-family: "SimHei", "黑体", sans-serif;\n'
        '    color: ' + Color.PRIMARY + ';\n'
        '    margin-bottom: 12pt;\n'
        '  }\n'
        '  .window {\n'
        '    text-align: center;\n'
        '    font-size: 10.5pt;\n'
        '    font-family: "仿宋", "FangSong", "SimSun", serif;\n'
        '    color: ' + Color.GRAY + ';\n'
        '    margin-bottom: 8pt;\n'
        '  }\n'
        '  .empty-decl {\n'
        '    text-align: center;\n'
        '    font-size: 12pt;\n'
        '    font-family: "仿宋", "FangSong", "SimSun", serif;\n'
        '    color: ' + Color.PRIMARY + ';\n'
        '    margin: 12pt 0;\n'
        '  }\n'
        '  .section { margin-bottom: 12pt; }\n'
        '  .section-title {\n'
        '    font-size: 14pt;\n'
        '    font-weight: bold;\n'
        '    font-family: "SimHei", "黑体", sans-serif;\n'
        '    color: ' + Color.ACCENT + ';\n'
        '    background: ' + Color.SHADE + ';\n'
        '    padding: 6pt 12pt;\n'
        '    border-left: 4mm solid ' + Color.ACCENT + ';\n'
        '    margin: 16pt 0 8pt;\n'
        '  }\n'
        '  .sub-title {\n'
        '    font-size: 12pt;\n'
        '    font-weight: bold;\n'
        '    font-family: "SimHei", "黑体", sans-serif;\n'
        '    color: ' + Color.PRIMARY + ';\n'
        '    margin: 8pt 0 4pt;\n'
        '  }\n'
        '  .item { margin-bottom: 8pt; }\n'
        '  .item-header {\n'
        '    font-weight: bold;\n'
        '    font-family: "SimHei", "黑体", sans-serif;\n'
        '    color: ' + Color.PRIMARY + ';\n'
        '  }\n'
        '  .item-body {\n'
        '    text-align: justify;\n'
        '    line-height: 1.5;\n'
        '    text-indent: 2em;\n'
        '    font-family: "仿宋", "FangSong", "SimSun", serif;\n'
        '  }\n'
        '  .source-line {\n'
        '    font-size: 10.5pt;\n'
        '    margin-top: 2pt;\n'
        '    font-family: "仿宋", "FangSong", "Times New Roman", serif;\n'
        '  }\n'
        '  .source-arrow { color: ' + Color.ACCENT + '; }\n'
        '  .source-label { color: ' + Color.GRAY + '; }\n'
        '  .source-url { color: ' + Color.SECONDARY + '; }\n'
        '  .empty-note {\n'
        '    font-size: 12pt;\n'
        '    font-family: "仿宋", "FangSong", "SimSun", serif;\n'
        '    color: ' + Color.GRAY + ';\n'
        '    margin: 6pt 0;\n'
        '  }\n'
        '  .analysis-text {\n'
        '    text-align: justify;\n'
        '    line-height: 1.5;\n'
        '    text-indent: 2em;\n'
        '    font-family: "仿宋", "FangSong", "SimSun", serif;\n'
        '    margin-bottom: 6pt;\n'
        '  }\n'
        '  .signal-item {\n'
        '    margin-bottom: 6pt;\n'
        '    line-height: 1.5;\n'
        '    text-indent: 2em;\n'
        '  }\n'
        '  .signal-label {\n'
        '    font-weight: bold;\n'
        '    font-family: "SimHei", "黑体", sans-serif;\n'
        '    color: ' + Color.ACCENT + ';\n'
        '  }\n'
        '  .footer {\n'
        '    border-top: 1px solid ' + Color.SECONDARY + ';\n'
        '    padding-top: 8pt;\n'
        '    margin-top: 24pt;\n'
        '    text-align: center;\n'
        '    font-size: 9pt;\n'
        '    font-family: "仿宋", "FangSong", "SimSun", serif;\n'
        '    color: ' + Color.SECONDARY + ';\n'
        '  }\n'
        '</style>\n'
        '</head>\n'
        '<body>'
    )


def _html_top_banner() -> str:
    return '<p class="top-badge">' + TOP_BADGE + '</p>'


def _html_title(date_cn: str, wday: str) -> str:
    return (
        '<p class="main-title">' + DOC_TITLE + '</p>'
        '<p class="en-subtitle">' + EN_SUBTITLE + '</p>'
        '<p class="date-line">' + date_cn + '\u3000' + wday + '</p>'
    )


def _html_item(item: dict) -> str:
    pub_date = item.get("pub_date", "")
    source = item.get("source", "")
    url = item.get("url", "")
    body = _strip_body_lead(item.get("body", ""))
    full_name = item.get("source_full_name")

    source_display = source
    if full_name:
        source_display = source + "\uff08" + full_name + "\uff09"

    header = pub_date + "\uff0c" + source_display + "\u62a5\u9053\uff1a"

    return (
        '<div class="item">'
        '<span class="item-header">' + _esc(header) + '</span>'
        ' <span class="item-body">' + _esc(body) + '</span>'
        '<div class="source-line">'
        '<span class="source-arrow">\u25b8 </span>'
        '<span class="source-label">\u6765\u6e90: </span>'
        '<span class="source-url">' + _esc(url) + '</span>'
        '</div>'
        '</div>'
    )


def _html_footer(date_cn: str) -> str:
    return (
        '<div class="footer">'
        + FOOTER_PREFIX + ' \u00b7 ' + date_cn
        + '</div>'
    )


# ═══════════════════════════════════════════════════════════════
#  D-4 · 输出路径白名单 (v3.2 可靠性, 2026-06-20)
# ═══════════════════════════════════════════════════════════════
# 用户强制约束: 日报三件套唯一输出位置 temp/output/daily_<YYYYMMDD>/
# spec: docs/specs/2026-06-20-daily-report-reliability-design.md §D-4
import re as _re_path
_ALLOWED_OUTPUT_PREFIX = "temp/output/daily_"
_ALLOWED_OUTPUT_RE = _re_path.compile(r"^temp/output/daily_\d{8}$")


def _check_output_dir(output_dir: str) -> None:
    """D-4 白名单校验: 仅允许 temp/output/daily_<8位日期>/ (无尾斜杠)。
    其它路径 (含 output/、./、绝对路径、漂到根目录) -> 抛 ValueError。
    用途: render.py CLI 入口、fallback 模块、调度 prompt 统一调用。
    """
    if not output_dir:
        raise ValueError("output_dir 为空")
    # 兼容尾部斜杠与反斜杠
    norm = output_dir.rstrip("/").rstrip("\\").replace("\\", "/")
    if not _ALLOWED_OUTPUT_RE.match(norm):
        raise ValueError(
            f"output_dir {output_dir!r} 不在白名单下;"
            f"日报三件套强制约束到 {_ALLOWED_OUTPUT_PREFIX!r}<YYYYMMDD>/"
        )


# ═══════════════════════════════════════════════════════════════
#  CLI 入口
# ═══════════════════════════════════════════════════════════════
def main():
    import argparse, sys as _sys_alias
    # R8 (2026-06-20): --format 作为 --fmt 别名, 向后兼容旧调用 (sop Appendix A #20)
    _alias_argv = [_sys_alias.argv[0]]
    for _a in _sys_alias.argv[1:]:
        if _a == '--format':
            _alias_argv.append('--fmt')
        elif _a.startswith('--format='):
            _alias_argv.append('--fmt=' + _a[len('--format='):])
        else:
            _alias_argv.append(_a)
    _sys_alias.argv = _alias_argv
    parser = argparse.ArgumentParser(description="日报三件套渲染器")
    parser.add_argument("json_path", help="report_data.json 路径")
    parser.add_argument("--fmt", "--format", dest="fmt", choices=["md", "docx", "html", "all"], default="all",
                        help="输出格式 (默认 all); 也接受 --format 作别名")
    parser.add_argument("--output-dir", default=None,
                        help="输出目录 (默认: temp/output/daily_<YYYYMMDD>/, 用户强制约束 D-4)")
    args = parser.parse_args()
    # 用户约束(2026-06-20起, D-4 升级): 默认指向 temp/output/daily_<YYYYMMDD>/ (北京时间)
    if args.output_dir is None:
        from datetime import datetime, timezone, timedelta
        _bjt = timezone(timedelta(hours=8))
        args.output_dir = f"temp/output/daily_{datetime.now(_bjt).strftime('%Y%m%d')}"
    # D-4 白名单校验
    try:
        _check_output_dir(args.output_dir)
    except ValueError as e:
        print(f"[ERROR] {e}", file=sys.stderr)
        sys.exit(4)
    import os
    os.makedirs(args.output_dir, exist_ok=True)

    data = load_data(args.json_path)
    # H1 24h 窗口守门 (v3.2, SOP 坑#13): window 字符串不符区间 → exit 2 触发 fallback。
    # B 路手工构造 (manual_construction=True) 豁免, 与 validate E.4-07 对齐。
    if not data.get("manual_construction"):
        try:
            enforce_window(data)
        except ValueError as e:
            print(f"[ERROR] enforce_window 拒绝: {e}", file=sys.stderr)
            sys.exit(2)
    fmt = args.fmt
    results = []

    if fmt in ("md", "all"):
        p = render_md(data, args.output_dir)
        results.append(("MD", p))
    if fmt in ("docx", "all"):
        p = render_docx(data, args.output_dir)
        results.append(("DOCX", p))
    if fmt in ("html", "all"):
        p = render_html(data, args.output_dir)
        results.append(("HTML", p))

    for label, fp in results:
        if fp:
            sz = os.path.getsize(fp)
            print("  [" + label + "] " + fp + "  (" + f"{sz:,}" + " bytes)")


if __name__ == "__main__":
    main()
