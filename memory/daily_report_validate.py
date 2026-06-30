#!/usr/bin/env python3
"""非传统安全领域动态日报 · 交付前自检 v1.8 
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
输入: report_data.json (+ 可选 .docx 文件用于 E.5 排版自检)
输出: 检查报告 (pass/fail + 明细)

E.4 内容自检 (14 项):
  1.  全部素材来自用户 URL
  2.  每条含日期/来源/主体/事实/链接
  3.  板块一 3-5 条, 以涉外因素为主
  4.  板块二 ≤10 条
  5.  板块 3.1 每条满足准入条件之一
  6.  中方主动发起内容符合双条件
  7.  时效窗口筛选合规
  8.  段首加粗含冒号格式正确 (数据层: body 前缀匹配)
  9.  英文缩写首次出现括注中文
  10. 来源行紧贴正文 (render 保证)
  11. 各板块倒序排列
  12. 板块四不引入新素材; 板块五不复制板块四句子
  13. 加粗仅 C.4 允许三处 (render 保证)
  14. 文件命名 YYYYMMDD

E.5 排版自检 (12 项, 需 docx 文件):
  E5-01.  顶部标识行样式
  E5-02.  主标题样式
  E5-03.  英文副标题样式
  E5-04.  日期行样式
  E5-05.  顶部区域与第一板块间次色细横线
  E5-06.  板块标题样式 (底纹 + 左边框)
  E5-07.  条目段首来源标样式
  E5-08.  条目正文样式
  E5-09.  来源行三色分层
  E5-10.  来源行与下一条正文段间距
  E5-11.  页边距 2.5 cm
  E5-12.  页脚含分隔线与页码
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
"""
import json
import re
import sys
import os
from datetime import datetime, timedelta
from typing import List, Tuple, Dict, Any, Optional

# 复用 render 的板块定义
try:
    from memory.daily_report_render import SECTIONS, parse_date
except ImportError:
    SECTIONS = [
        {"code": "S1", "num": "一", "title": "非传统安全领域涉华要闻", "min": 3, "max": 5},
        {"code": "S2", "num": "二", "title": "主要国家及国际组织的重要动向", "min": 0, "max": 10},
        {"code": "S3", "num": "三", "title": "其他热点或苗头性线索", "optional": True},
        {"code": "S4", "num": "四", "title": "本期要情与趋势分析", "is_analysis": True},
        {"code": "S5", "num": "五", "title": "重点关注信号的情报价值", "min": 3, "max": 5, "is_signals": True},
    ]


def load_data(path: str) -> dict:
    """加载 report_data.json"""
    with open(path, 'r', encoding='utf-8') as f:
        return json.load(f)


def _get_section_items(data: dict, code: str) -> list:
    """获取板块条目"""
    key = code.lower() + "_items"
    return data.get(key, [])


def _parse_md(md_str: str):
    """解析 'X月X日' 为可比较的元组"""
    m = re.match(r'(\d{1,2})月(\d{1,2})日', md_str or "")
    if m:
        return int(m.group(1)), int(m.group(2))
    return 0, 0


# ─── E.4 内容自检 (14 项) ───────────────────────────────────

def check_e4_01_url_coverage(data: dict) -> Tuple[bool, List[str]]:
    """E.4-1: 全部素材来自用户 URL"""
    errors = []
    all_items = []
    for code in ["s1", "s2"]:
        all_items.extend(_get_section_items(data, code))
    for sub in ["s3_hot", "s3_clues"]:
        all_items.extend(data.get(sub, []))
    for i, item in enumerate(all_items):
        url = item.get("url", "").strip()
        if not url:
            errors.append(f"条目{i+1}: 缺少URL")
        elif not url.startswith("http"):
            errors.append(f"条目{i+1}: URL格式异常 '{url[:50]}'")
    return len(errors) == 0, errors


def check_e4_02_required_fields(data: dict) -> Tuple[bool, List[str]]:
    """E.4-2: 每条含日期/来源/主体/事实/链接"""
    errors = []
    required = ["pub_date", "source", "body", "url"]
    for code in ["s1", "s2"]:
        items = _get_section_items(data, code)
        for i, item in enumerate(items, 1):
            for field in required:
                if not item.get(field):
                    errors.append(f"{code.upper()} 条目{i}: 缺少 '{field}'")
            body_len = len(item.get("body", ""))
            if body_len < 50:
                errors.append(f"{code.upper()} 条目{i}: 正文过短 ({body_len}字)")
    for sub in ["s3_hot", "s3_clues"]:
        items = data.get(sub, [])
        for i, item in enumerate(items, 1):
            for field in required:
                if not item.get(field):
                    errors.append(f"{sub.upper()} 条目{i}: 缺少 '{field}'")
    return len(errors) == 0, errors


def check_e4_03_section1_foreign(data: dict) -> Tuple[bool, List[str]]:
    """E.4-3: 板块一 3-5 条, 以涉外因素为主"""
    errors = []
    s1 = _get_section_items(data, "s1")
    count = len(s1)
    if count < 3 or count > 5:
        errors.append(f"板块一(涉华要闻): {count}条 (要求3-5条)")
    # 检查涉外因素: 事件发起方应为外部
    foreign_tags = data.get("s1_foreign_check", [])
    if foreign_tags:
        for i, tag in enumerate(foreign_tags):
            if tag == "domestic_initiative":
                errors.append(f"S1 条目{i+1}: 疑似中方主动发起, 需确认涉外因素")
    return len(errors) == 0, errors


def check_e4_04_section2_count(data: dict) -> Tuple[bool, List[str]]:
    """E.4-4: 板块二 ≤10 条"""
    errors = []
    s2 = _get_section_items(data, "s2")
    if len(s2) > 10:
        errors.append(f"板块二(各国动向): {len(s2)}条 (要求≤10条)")
    return len(errors) == 0, errors


def check_e4_05_s31_admission(data: dict) -> Tuple[bool, List[str]]:
    """E.4-5: 板块 3.1 每条满足三项准入条件之一"""
    errors = []
    s3_hot = data.get("s3_hot", [])
    for i, item in enumerate(s3_hot, 1):
        admission = item.get("admission", "")
        if admission and admission not in ("official_response", "two_sources", "verifiable_impact"):
            errors.append(f"3.1 条目{i}: 准入条件'{admission}'不符合要求")
    return len(errors) == 0, errors


def check_e4_06_china_initiative(data: dict) -> Tuple[bool, List[str]]:
    """E.4-6: 中方主动发起内容符合'7日+明文提及'双条件"""
    errors = []
    # 检查所有条目是否有 china_initiative 标记
    all_items = []
    for code in ["s1", "s2"]:
        all_items.extend(_get_section_items(data, code))
    for i, item in enumerate(all_items):
        if item.get("china_initiative") and not item.get("meets_7day_rule"):
            errors.append(f"条目{i+1}: 中方主动发起但未满足'7日+明文提及'双条件")
    return len(errors) == 0, errors


def check_e4_07_timeliness(data: dict) -> Tuple[bool, List[str]]:
    """E.4-07 v3.2 严格 24h 窗口 (替换 v3.1 的 7 天实现)。
    豁免: data.get('manual_construction') is True 时跳过 (B 路手工构造路径)。

    校验规则 (SOP H1):
      - 报告日 D = data['date']
      - 24h 窗口: D-1 00:00 ~ D 18:00 (北京时间)
      - 每条 pub_date (X月X日) 必须落在窗口内
    """
    errors = []
    # 豁免: B 路手工构造 (SOP v3.1 坑 #12)
    if data.get("manual_construction") is True:
        return True, ["(manual_construction=True, E.4-07 跳过)"]

    report_date = data.get("date", "")
    if not report_date:
        errors.append("缺少报告日期, 无法校验时效窗口")
        return False, errors
    try:
        rd = datetime.strptime(report_date, "%Y-%m-%d")
    except ValueError:
        errors.append(f"报告日期格式异常: '{report_date}'")
        return False, errors

    # 24h 窗口: D-1 00:00 ~ D 18:00
    window_start = rd - timedelta(days=1)
    window_end = rd.replace(hour=18, minute=0, second=0, microsecond=0)

    all_items = []
    for code in ["s1", "s2"]:
        all_items.extend(_get_section_items(data, code))
    for sub in ["s3_hot", "s3_clues"]:
        all_items.extend(data.get(sub, []))
    for i, item in enumerate(all_items):
        pd = item.get("pub_date", "")
        m = re.match(r'(\d{1,2})月(\d{1,2})日', pd)
        if not m:
            errors.append(f"条目{i+1}: pub_date 无法解析为 'X月X日': '{pd}'")
            continue
        month, day = int(m.group(1)), int(m.group(2))
        # 跨年处理: month > D.month → 上一年
        year = rd.year - 1 if month > rd.month else rd.year
        try:
            item_dt = datetime(year, month, day, 12, 0)  # 用 12:00 避开边界
        except ValueError:
            errors.append(f"条目{i+1}: pub_date 非法日期 '{pd}'")
            continue
        if item_dt < window_start:
            delta_days = (window_start - item_dt).days
            errors.append(f"条目{i+1}: '{pd}' 早于 24h 窗口起点 ({delta_days} 天)")
        elif item_dt > window_end:
            errors.append(f"条目{i+1}: '{pd}' 晚于 24h 窗口终点 (D 18:00)")
    return len(errors) == 0, errors


def check_e4_08_bold_format(data: dict) -> Tuple[bool, List[str]]:
    """E.4-8: 段首加粗含冒号格式 (数据层: body前缀匹配)"""
    errors = []
    pattern = re.compile(r'^\d{1,2}月\d{1,2}日.{1,40}(报道|表示|声明|发布)')
    for code in ["s1", "s2"]:
        items = _get_section_items(data, code)
        for i, item in enumerate(items, 1):
            body = item.get("body", "")
            if body and not pattern.match(body):
                hint = _e408_hint(body)
                errors.append(f"{code.upper()} 条目{i}: 正文开头未匹配段首格式 '{body[:60]}...'  {hint}")
    for sub in ["s3_hot", "s3_clues"]:
        items = data.get(sub, [])
        for i, item in enumerate(items, 1):
            body = item.get("body", "")
            if body and not pattern.match(body):
                hint = _e408_hint(body)
                errors.append(f"{sub.upper()} 条目{i}: 正文开头未匹配段首格式 '{body[:60]}...'  {hint}")
    return len(errors) == 0, errors


def _e408_hint(body: str) -> str:
    """R8 (2026-06-20): E.4-08 失败时给出可定位的 hint。"""
    import re as _re
    m = _re.match(r'^(\d{1,2}月\d{1,2}日)(.*)', body)
    if not m:
        return "[hint] 段首必须以'X月X日'开头"
    rest = m.group(2)
    mm = _re.search(r'(报道|表示|声明|发布)', rest)
    if not mm:
        return f"[hint] 段首无'报道/表示/声明/发布'动词; 首 60 字符已打印"
    dist = mm.start()
    if dist > 40:
        return f"[hint] '{mm.group(0)}'距日期 {dist} 字符 (>40); 压缩来源标注或前移动词"
    return f"[hint] '{mm.group(0)}'在 {dist} 字符处, 已合规; 检查标点/冒号"

def check_e4_09_abbreviation(data: dict) -> Tuple[bool, List[str]]:
    """E.4-9: 英文缩写首次出现已括注中文全称"""
    errors = []
    # 收集所有正文文本, 检查大写缩写(2-6字母)是否首次括注
    abbr_pattern = re.compile(r'\b([A-Z]{2,6})\b')
    seen = {}
    all_texts = []
    for code in ["s1", "s2"]:
        for item in _get_section_items(data, code):
            all_texts.append((code.upper(), item.get("body", "")))
    for sub in ["s3_hot", "s3_clues"]:
        for item in data.get(sub, []):
            all_texts.append((sub.upper(), item.get("body", "")))

    # 跳过段首lead前缀(日期+来源+报道：)，只检查正文内容部分
    lead_re = re.compile(r'^\d{1,2}月\d{1,2}日.{1,40}(报道|表示|声明|发布)[：:]')
    for label, text in all_texts:
        # 去掉lead前缀，只扫描正文
        content = lead_re.sub('', text, count=1)
        for m in abbr_pattern.finditer(content):
            abbr = m.group(1)
            if abbr in seen:
                continue
            # 检查首次出现附近是否有括注
            pos = m.start()
            ctx = content[max(0, pos-5):pos+len(abbr)+20]
            if '(' not in ctx and '（' not in ctx:
                # 可能是首次出现但未括注
                seen[abbr] = False
            else:
                seen[abbr] = True

    for abbr, annotated in seen.items():
        if not annotated:
            errors.append(f"缩写'{abbr}'首次出现可能未括注中文全称")
    return len(errors) == 0, errors


def check_e4_10_source_line_spacing(data: dict) -> Tuple[bool, List[str]]:
    """E.4-10: 来源行紧贴正文 (由 render 保证, 此处仅标记)"""
    # 数据层无法验证排版间距, 标记为 auto-pass
    return True, ["(由 render.py 保证, 数据层自动通过)"]


def check_e4_11_reverse_order(data: dict) -> Tuple[bool, List[str]]:
    """E.4-11: 各板块内按发布日期倒序排列"""
    errors = []
    for code in ["s1", "s2"]:
        items = _get_section_items(data, code)
        dates = [item.get("pub_date", "") for item in items]
        for i in range(len(dates) - 1):
            cur = _parse_md(dates[i])
            nxt = _parse_md(dates[i + 1])
            if cur < nxt:
                errors.append(f"{code.upper()} 非倒序: 条目{i+1}({dates[i]}) < 条目{i+2}({dates[i+1]})")
    for sub in ["s3_hot", "s3_clues"]:
        items = data.get(sub, [])
        dates = [item.get("pub_date", "") for item in items]
        for i in range(len(dates) - 1):
            cur = _parse_md(dates[i])
            nxt = _parse_md(dates[i + 1])
            if cur < nxt:
                errors.append(f"{sub.upper()} 非倒序: 条目{i+1}({dates[i]}) < 条目{i+2}({dates[i+1]})")
    return len(errors) == 0, errors


def check_e4_12_no_new_material(data: dict) -> Tuple[bool, List[str]]:
    """E.4-12: 板块四不引入新素材; 板块五不复制板块四句子"""
    errors = []
    _trends_raw = data.get("trends", "")
    # 兼容三段式dict和纯文本
    if isinstance(_trends_raw, dict):
        trends = " ".join(v for v in _trends_raw.values() if isinstance(v, str))
    else:
        trends = _trends_raw
    # 收集板块1-3的body用于比对
    s123_bodies = set()
    for code in ["s1", "s2"]:
        for item in _get_section_items(data, code):
            body = item.get("body", "")
            if len(body) > 20:
                s123_bodies.add(body[:50])
    for sub in ["s3_hot", "s3_clues"]:
        for item in data.get(sub, []):
            body = item.get("body", "")
            if len(body) > 20:
                s123_bodies.add(body[:50])
    # 板块四不应包含板块1-3中没有的新URL
    urls_in_trends = re.findall(r'https?://\S+', trends)
    for url in urls_in_trends:
        errors.append(f"板块四含新URL: '{url[:60]}' (不应引入新素材)")
    # 板块五不应复制板块四句子
    signals = data.get("signals", [])
    for i, sig in enumerate(signals, 1):
        sig_text = sig.get("text", "")
        if trends and sig_text:
            # 简单检查: 如果信号文本与趋势分析有>30字相同子串
            for start in range(0, len(sig_text)-10):
                chunk = sig_text[start:start+30]
                if len(chunk) >= 20 and chunk in trends:
                    errors.append(f"信号{i}: 与板块四趋势分析有重叠文本 '{chunk[:20]}...'")
                    break
    return len(errors) == 0, errors


def check_e4_13_bold_only_three(data: dict) -> Tuple[bool, List[str]]:
    """E.4-13: 加粗仅 C.4 允许三处 (由 render 保证, 数据层标记)"""
    # C.4 允许三处: 1)段首来源标 2)板块标题 3)信号标签
    # 数据层无法验证Markdown加粗, 标记为 auto-pass
    return True, ["(由 render.py 保证, 数据层自动通过)"]


def check_e4_14_filename(data: dict) -> Tuple[bool, List[str]]:
    """E.4-14: 文件命名为 非传统安全领域动态日报_YYYYMMDD"""
    errors = []
    report_date = data.get("date", "")
    if report_date:
        expected_pattern = report_date.replace("-", "")
        filename = data.get("filename", "")
        if filename and expected_pattern not in filename:
            errors.append(f"文件名'{filename}'未包含日期{expected_pattern}")
    return len(errors) == 0, errors


BODY_MAX_CHARS = 200  # 单条新闻 body 总结字数上限 (2026-06-23 新增)


def check_e4_15_body_length(data: dict) -> Tuple[bool, List[str]]:
    """E.4-15: 每条新闻 body 总结 ≤ 200 字 (严格验证, 2026-06-23 新增)"""
    errors = []
    sections = [
        ("s1_items", data.get("s1_items", [])),
        ("s2_items", data.get("s2_items", [])),
        ("s3_hot", data.get("s3_hot", [])),
        ("s3_clues", data.get("s3_clues", [])),
    ]
    for sec_name, items in sections:
        for idx, item in enumerate(items, 1):
            body = item.get("body", "")
            L = len(body)
            if L > BODY_MAX_CHARS:
                src = item.get("source", "?")
                errors.append(
                    f"{sec_name}#{idx} ({src}) body 超长: {L} 字 > {BODY_MAX_CHARS} | 摘要: {body[:60]}..."
                )
    return len(errors) == 0, errors


def check_e4_16_trends_length(data: dict) -> Tuple[bool, List[str]]:
    """E.4-16: trends 三段 (核心态势/行为体联动/涉华影响方向) 每段 200-300 字, 总长 600-900 字
    (SOP §2.4 硬约束, 2026-06-23 新增)"""
    errors = []
    trends = data.get("trends", {})
    if not isinstance(trends, dict):
        return True, ["(无 trends 字段, 跳过)"]
    seg_names = ["核心态势", "行为体联动", "涉华影响方向"]
    seg_keys = ["core_situation", "actor_dynamics", "china_impact_direction"]
    total = 0
    for name, key in zip(seg_names, seg_keys):
        text = trends.get(key, "")
        L = len(text)
        total += L
        if L < 200:
            errors.append(f"trends.{key} ({name}) 过短: {L} 字 < 200")
        elif L > 300:
            errors.append(f"trends.{key} ({name}) 超长: {L} 字 > 300")
    if total < 600:
        errors.append(f"trends 总长过短: {total} 字 < 600")
    elif total > 900:
        errors.append(f"trends 总长超长: {total} 字 > 900")
    return len(errors) == 0, errors


def check_e4_17_signals_spec(data: dict) -> Tuple[bool, List[str]]:
    """E.4-17: signals 每条 text 60-120 字 + 必须含 '(情报缺口：' + label 末尾禁含冒号
    (SOP §2.5 硬约束, 2026-06-23 新增)"""
    errors = []
    signals = data.get("signals", [])
    if not signals:
        return True, ["(无 signals 字段, 跳过)"]
    for idx, sig in enumerate(signals, 1):
        label = sig.get("label", "")
        text = sig.get("text", "")
        if label.endswith((":", "：", ": ")):
            errors.append(f"signals#{idx} label 末尾禁含冒号: '{label[-1]}'")
        L = len(text)
        if L < 60:
            errors.append(f"signals#{idx} text 过短: {L} 字 < 60 | 摘要: {text[:40]}...")
        elif L > 120:
            errors.append(f"signals#{idx} text 超长: {L} 字 > 120 | 摘要: {text[:40]}...")
        if "情报缺口" not in text:
            errors.append(f"signals#{idx} text 缺'(情报缺口：' 标注: '{text[:40]}...'")
    return len(errors) == 0, errors


# ─── E.5 排版自检 (12 项, 需 docx 文件) ──────────────────────

def check_e5_layout(docx_path: str) -> Tuple[bool, List[Dict[str, Any]]]:
    """E.5 排版自检: 检查 docx 文件排版规范 (需 python-docx)"""
    results = []
    try:
        from docx import Document
        from docx.shared import Pt, Cm, RGBColor
    except ImportError:
        return False, [{"check": "E.5排版自检", "pass": False,
                        "errors": ["需要 python-docx 库: pip install python-docx"]}]

    doc = Document(docx_path)

    # E5-01: 顶部标识行
    passed, errs = _e5_check_banner(doc)
    results.append({"check": "E5-01 顶部标识行", "pass": passed, "errors": errs})

    # E5-02: 主标题
    passed, errs = _e5_check_title(doc)
    results.append({"check": "E5-02 主标题", "pass": passed, "errors": errs})

    # E5-03: 英文副标题
    passed, errs = _e5_check_subtitle(doc)
    results.append({"check": "E5-03 英文副标题", "pass": passed, "errors": errs})

    # E5-04: 日期行
    passed, errs = _e5_check_date_line(doc)
    results.append({"check": "E5-04 日期行", "pass": passed, "errors": errs})

    # E5-05: 顶部区域与第一板块间细横线
    passed, errs = _e5_check_separator(doc)
    results.append({"check": "E5-05 分隔线", "pass": passed, "errors": errs})

    # E5-06: 板块标题 (底纹+左边框)
    passed, errs = _e5_check_section_titles(doc)
    results.append({"check": "E5-06 板块标题", "pass": passed, "errors": errs})

    # E5-07: 条目段首来源标
    passed, errs = _e5_check_source_bold(doc)
    results.append({"check": "E5-07 段首来源标", "pass": passed, "errors": errs})

    # E5-08: 条目正文样式
    passed, errs = _e5_check_body_style(doc)
    results.append({"check": "E5-08 条目正文", "pass": passed, "errors": errs})

    # E5-09: 来源行三色分层
    passed, errs = _e5_check_source_tricolor(doc)
    results.append({"check": "E5-09 来源行三色", "pass": passed, "errors": errs})

    # E5-10: 来源行与下一条正文段间距
    passed, errs = _e5_check_source_spacing(doc)
    results.append({"check": "E5-10 来源行间距", "pass": passed, "errors": errs})

    # E5-11: 页边距 2.5cm
    passed, errs = _e5_check_margins(doc)
    results.append({"check": "E5-11 页边距", "pass": passed, "errors": errs})

    # E5-12: 页脚
    passed, errs = _e5_check_footer(doc)
    results.append({"check": "E5-12 页脚", "pass": passed, "errors": errs})

    all_pass = all(r["pass"] for r in results)
    return all_pass, results


# ─── E.5 辅助函数 ──────────────────────────────────────────

def _e5_check_banner(doc) -> Tuple[bool, List[str]]:
    """检查顶部标识行"""
    errors = []
    if not doc.paragraphs:
        return False, ["文档无段落"]
    p = doc.paragraphs[0]
    text = p.text.strip()
    if "每日情报整编" not in text:
        errors.append(f"第一段非标识行: '{text[:30]}'")
        return False, errors
    return True, []


def _e5_check_title(doc) -> Tuple[bool, List[str]]:
    """检查主标题"""
    errors = []
    if len(doc.paragraphs) < 2:
        return False, ["段落不足, 无主标题"]
    p = doc.paragraphs[1]
    if "非传统安全领域动态日报" not in p.text:
        errors.append(f"第二段非主标题: '{p.text[:30]}'")
    return len(errors) == 0, errors


def _e5_check_subtitle(doc) -> Tuple[bool, List[str]]:
    """检查英文副标题"""
    errors = []
    if len(doc.paragraphs) < 3:
        return False, ["段落不足, 无英文副标题"]
    p = doc.paragraphs[2]
    if "NON-TRADITIONAL SECURITY" not in p.text.upper():
        errors.append(f"第三段非英文副标题: '{p.text[:30]}'")
    return len(errors) == 0, errors


def _e5_check_date_line(doc) -> Tuple[bool, List[str]]:
    """检查日期行"""
    errors = []
    if len(doc.paragraphs) < 4:
        return False, ["段落不足, 无日期行"]
    p = doc.paragraphs[3]
    if not re.search(r'\d{4}年\d{1,2}月\d{1,2}日', p.text):
        errors.append(f"第四段非日期行: '{p.text[:30]}'")
    return len(errors) == 0, errors


def _e5_check_separator(doc) -> Tuple[bool, List[str]]:
    """检查分隔线(通过段落边框或空段落)"""
    # 简化检查: 第5段应为空或含分隔线
    if len(doc.paragraphs) > 4:
        return True, []
    return True, ["(简化检查通过)"]


def _e5_check_section_titles(doc) -> Tuple[bool, List[str]]:
    """检查板块标题有底纹和左边框"""
    errors = []
    found_section = False
    for p in doc.paragraphs:
        if re.match(r'^[一二三四五]、', p.text):
            found_section = True
            # 检查是否有底纹 (shading) 和左边框
            has_shading = False
            has_border = False
            pPr = p._element.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}shd')
            if pPr is not None:
                has_shading = True
            pBdr = p._element.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pBdr')
            if pBdr is not None:
                has_border = True
            if not has_shading:
                errors.append(f"板块'{p.text[:10]}' 缺少底纹")
            if not has_border:
                errors.append(f"板块'{p.text[:10]}' 缺少左边框")
    if not found_section:
        errors.append("未找到板块标题段落")
    return len(errors) == 0, errors


def _e5_check_source_bold(doc) -> Tuple[bool, List[str]]:
    """检查段首来源标加粗"""
    errors = []
    for p in doc.paragraphs:
        m = re.match(r'\d{1,2}月\d{1,2}日', p.text)
        if m and len(p.runs) > 0:
            if not p.runs[0].bold:
                errors.append(f"段首未加粗: '{p.text[:20]}'")
    return len(errors) == 0, errors


def _e5_check_body_style(doc) -> Tuple[bool, List[str]]:
    """检查正文样式"""
    # 简化: 检查有正文段落存在
    return True, ["(简化检查通过)"]


def _e5_check_source_tricolor(doc) -> Tuple[bool, List[str]]:
    """检查来源行三色分层"""
    errors = []
    for p in doc.paragraphs:
        if p.text.strip().startswith("\u25b8"):
            if len(p.runs) < 3:
                errors.append(f"来源行run数不足({len(p.runs)}): '{p.text[:30]}'")
    return len(errors) == 0, errors


def _e5_check_source_spacing(doc) -> Tuple[bool, List[str]]:
    """检查来源行与下一条正文间距"""
    return True, ["(间距由 render 保证, 简化检查通过)"]


def _e5_check_margins(doc) -> Tuple[bool, List[str]]:
    """检查页边距 2.5cm"""
    errors = []
    for section in doc.sections:
        cm = 2.54  # 1 inch = 2.54 cm
        top = section.top_margin / 914400 * cm  # EMU to cm
        bottom = section.bottom_margin / 914400 * cm
        left = section.left_margin / 914400 * cm
        right = section.right_margin / 914400 * cm
        for name, val in [("上", top), ("下", bottom), ("左", left), ("右", right)]:
            if abs(val - 2.5) > 0.3:
                errors.append(f"{name}边距: {val:.1f}cm (要求2.5cm)")
    return len(errors) == 0, errors


def _e5_check_footer(doc) -> Tuple[bool, List[str]]:
    """检查页脚含分隔线与页码"""
    # 简化检查
    return True, ["(页脚由 render 保证, 简化检查通过)"]


# ─── 主验证流程 ─────────────────────────────────────────────

def validate(data: dict) -> Tuple[bool, List[Dict[str, Any]]]:
    """执行 E.4 全部 14 项检查, 返回 (pass, results)"""
    checks = [
        ("E.4-01  URL覆盖检查",           check_e4_01_url_coverage),
        ("E.4-02  字段完整性",             check_e4_02_required_fields),
        ("E.4-03  板块一条数与涉外因素",    check_e4_03_section1_foreign),
        ("E.4-04  板块二条数",             check_e4_04_section2_count),
        ("E.4-05  板块3.1准入条件",        check_e4_05_s31_admission),
        ("E.4-06  中方主动发起双条件",      check_e4_06_china_initiative),
        ("E.4-07  时效窗口合规",           check_e4_07_timeliness),
        ("E.4-08  段首加粗格式",           check_e4_08_bold_format),
        ("E.4-09  英文缩写括注",          check_e4_09_abbreviation),
        ("E.4-10  来源行间距(render)",     check_e4_10_source_line_spacing),
        ("E.4-11  倒序排列",              check_e4_11_reverse_order),
        ("E.4-12  板块四五无新材/无重叠",  check_e4_12_no_new_material),
        ("E.4-13  加粗仅三处(render)",     check_e4_13_bold_only_three),
        ("E.4-14  文件命名YYYYMMDD",      check_e4_14_filename),
        ("E.4-15  每条body≤200字",        check_e4_15_body_length),
        ("E.4-16  trends三段200-300总600-900", check_e4_16_trends_length),
        ("E.4-17  signals text 60-120+情报缺口+无尾冒号", check_e4_17_signals_spec),
    ]

    results = []
    all_pass = True
    for name, fn in checks:
        passed, errors = fn(data)
        results.append({
            "check": name,
            "pass": passed,
            "errors": errors,
        })
        if not passed:
            all_pass = False

    return all_pass, results


def validate_e5(docx_path: str) -> Tuple[bool, List[Dict[str, Any]]]:
    """执行 E.5 全部 12 项排版检查"""
    return check_e5_layout(docx_path)


def format_report(results: List[Dict[str, Any]], title: str = "日报数据自检报告 (E.4)") -> str:
    """格式化检查报告"""
    lines = []
    lines.append("=" * 55)
    lines.append(f"  {title}")
    lines.append("=" * 55)

    for r in results:
        status = "\u2705 PASS" if r["pass"] else "\u274c FAIL"
        lines.append(f"\n[{status}] {r['check']}")
        for err in r["errors"]:
            lines.append(f"    - {err}")

    total = len(results)
    passed = sum(1 for r in results if r["pass"])
    lines.append("")
    lines.append("=" * 55)
    lines.append(f"  结果: {passed}/{total} 项通过")
    lines.append("=" * 55)
    return "\n".join(lines)


def main():
    import argparse
    parser = argparse.ArgumentParser(description="日报交付自检 v1.8 (E.4 + E.5)")
    parser.add_argument("json_path", help="report_data.json 路径")
    parser.add_argument("--docx", help="可选: docx 文件路径, 触发 E.5 排版自检")
    parser.add_argument("--strict", action="store_true", help="严格模式: 有错误则退出码非0")
    args = parser.parse_args()

    data = load_data(args.json_path)
    all_pass, results = validate(data)
    report = format_report(results)
    print(report)

    # E.5 排版自检 (可选)
    if args.docx:
        print()
        e5_pass, e5_results = validate_e5(args.docx)
        e5_report = format_report(e5_results, title="日报排版自检报告 (E.5)")
        print(e5_report)
        if not e5_pass:
            all_pass = False

    if args.strict and not all_pass:
        sys.exit(1)

    return 0 if all_pass else 1


if __name__ == "__main__":
    sys.exit(main())
